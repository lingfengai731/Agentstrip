from __future__ import annotations

from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageOps


ROOT = Path(__file__).resolve().parents[1]
TEAL = "0B4D43"
GOLD = "D5A62A"
PAPER = "F7F2E7"
PALE_TEAL = "E8F2EF"
INK = "172A2A"
MUTED = "5E6A68"


DRIVERS = {
    "dicky": {
        "name": "Dicky",
        "pack": ROOT / "wandermind-studio/promotion-packs/Dicky",
        "link": "https://wandermind.cc/find-driver?driver_id=dicky",
        "utm": "https://wandermind.cc/find-driver?driver_id=dicky&utm_source=driver_dicky&utm_medium=instagram&utm_campaign=bali_driver_launch",
        "service": "机场接送、巴厘岛私人一日游，以及根据客人行程方向安排顺路路线。",
        "photos": [
            ("01-Dicky-profile.jpg", "01 · 个人照片", "可作为首图；如果不想发个人照片，也可从本资料包的风景图开始。"),
            ("02-Dicky-guest-moment.jpg", "02 · 与客人的真实服务时刻", "体现真实服务经验，但公开文案中不要写出客人身份。"),
            ("03-Dicky-vehicle.jpg", "03 · 车辆参考", "仅作车辆视觉参考；具体车辆、座位和行李容量仍须逐单确认。"),
            ("04-Dicky-Melasti-Beach.jpg", "04 · Melasti Beach", "Dare2Leap · CC BY-SA 4.0"),
            ("05-Dicky-Jimbaran-Sunset.jpg", "05 · Jimbaran Bay", "Simon Sees · CC BY 2.0"),
            ("06-Dicky-Broken-Beach.jpg", "06 · Broken Beach", "Aaron Rentfrew · CC BY-SA 4.0"),
            ("07-Dicky-Tegallalang.jpg", "07 · Tegallalang", "Philip Nalangan · CC BY 4.0"),
        ],
        "scenic_caption": "一次旅程看见巴厘岛的四种面貌：Melasti 的南部海岸、Jimbaran 的日落、Broken Beach 的天然海拱，以及 Tegallalang 的层叠稻田。先收藏灵感，再打开我的 WanderMind 专属链接提交计划。路线、日期、车辆、时长和最终价格，都将在收到请求后确认。",
    },
    "gede": {
        "name": "Gede Nico",
        "pack": ROOT / "wandermind-studio/promotion-packs/Gede-Nico",
        "link": "https://wandermind.cc/find-driver?driver_id=gede",
        "utm": "https://wandermind.cc/find-driver?driver_id=gede&utm_source=driver_gede&utm_medium=instagram&utm_campaign=bali_driver_launch",
        "service": "把巴厘岛文化、美食、活动和本地体验串成更顺路、可执行的一天。",
        "photos": [
            ("01-Gede-profile.jpg", "01 · 个人照片", "可作为首图；如果不想发个人照片，也可从本资料包的风景图开始。"),
            ("02-Gede-guest-moment.jpg", "02 · 与客人的真实服务时刻", "体现真实服务经验，但公开文案中不要写出客人身份。"),
            ("03-Gede-vehicle.jpg", "03 · 车辆参考", "仅作车辆视觉参考；具体车辆、座位和行李容量仍须逐单确认。"),
            ("04-Gede-Campuhan-Ridge.jpg", "04 · Campuhan Ridge Walk", "Artem Beliaikin · CC0 1.0"),
            ("05-Gede-Jatiluwih.jpg", "05 · Jatiluwih", "Jorge Franganillo · CC BY 2.0"),
            ("06-Gede-Tirta-Gangga.jpg", "06 · Tirta Gangga", "Bair175 · CC BY-SA 3.0"),
            ("07-Gede-Seminyak-Sunset.jpg", "07 · Seminyak Beach", "Christophe95 · CC BY-SA 4.0"),
        ],
        "scenic_caption": "巴厘岛的四种氛围：Campuhan 的绿色清晨、Jatiluwih 的层叠稻田、Tirta Gangga 的水上花园，以及 Seminyak 的日落。先收藏灵感，再打开我的 WanderMind 专属链接提交计划。路线、日期、车辆、时长和最终价格，都将在收到请求后确认。",
    },
}


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.4)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, color, before, after in (
        ("Title", 28, TEAL, 0, 8),
        ("Subtitle", 12, MUTED, 0, 12),
        ("Heading 1", 18, TEAL, 16, 8),
        ("Heading 2", 13, TEAL, 11, 5),
        ("Heading 3", 11, GOLD, 8, 3),
    ):
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = style_name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(max(after - 1, 2))
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("WanderMind Studio · 中文审阅版")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(TEAL)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("内部审阅材料 · 未经确认不要直接对外发布")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def add_brand_rule(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_cell_width(table.cell(0, 0), 6500)
    set_cell_width(table.cell(0, 1), 2860)
    shade(table.cell(0, 0), TEAL)
    shade(table.cell(0, 1), GOLD)
    for cell in table.rows[0].cells:
        cell.height = Pt(4)
        cell_margins(cell, 0, 0, 0, 0)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)


def add_callout(doc: Document, label: str, text: str, fill=PAPER) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_cell_width(table.cell(0, 0), 9360)
    cell = table.cell(0, 0)
    shade(cell, fill)
    cell_margins(cell, 150, 170, 150, 170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}  ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)
    p.add_run(text)


def add_copy_box(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for cell in (table.cell(0, 0), table.cell(1, 0)):
        set_cell_width(cell, 9360)
        cell_margins(cell, 75, 140, 75, 140)
    shade(table.cell(0, 0), TEAL)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    shade(table.cell(1, 0), PAPER)
    p = table.cell(1, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.add_run(text)


def add_photo_card(cell, image_path: Path, title: str, note: str, width=Inches(2.85)) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cell_margins(cell, 100, 100, 100, 100)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    p.add_run().add_picture(image_stream(image_path), width=width)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)
    p = cell.add_paragraph(note)
    p.paragraph_format.space_after = Pt(0)
    p.runs[0].font.size = Pt(8.5)
    p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)


def image_stream(image_path: Path) -> BytesIO:
    """Return a compact metadata-free JPEG stream without modifying the source file."""
    with Image.open(image_path) as source:
        image = ImageOps.exif_transpose(source).convert("RGB")
        image.thumbnail((1800, 1800), Image.Resampling.LANCZOS)
        stream = BytesIO()
        image.save(stream, format="JPEG", quality=84, optimize=True, progressive=True)
    stream.seek(0)
    return stream


def add_cover(doc: Document, d: dict) -> None:
    add_brand_rule(doc)
    p = doc.add_paragraph("WanderMind Studio", style="Subtitle")
    p.paragraph_format.space_before = Pt(18)
    p = doc.add_paragraph("完整合作与推广指南", style="Title")
    p = doc.add_paragraph(d["name"], style="Title")
    p.runs[0].font.color.rgb = RGBColor.from_string(GOLD)
    doc.add_paragraph("网站介绍 · 手机使用 · 两组推广素材 · 报价授权", style="Subtitle")

    hero = d["pack"] / d["photos"][0][0]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(image_stream(hero), width=Inches(4.35))
    add_callout(doc, "审阅说明", "本文件用于你先检查中文含义。司机实际收到和使用的仍是各自的印尼语版本；姓名、专属链接和图片包不得互换。")

    doc.add_paragraph("7 张可用图片 · 2 组现成发布 · 1 条司机专属链接", style="Subtitle")


def add_site_orientation(doc: Document, d: dict) -> None:
    doc.add_page_break()
    doc.add_paragraph("先了解 WanderMind", style="Heading 1")
    add_callout(doc, "我们的初衷", "WanderMind 希望让旅行者先看懂真实地点、路线强度与时间成本，再把完整需求交给本地司机确认。网站不把精修照片当作现实，也不把未核验价格写成最终报价。", PALE_TEAL)
    doc.add_paragraph("网站能帮助游客做什么", style="Heading 2")
    for item in (
        "浏览巴厘岛公开路线、地点实景和可编辑的一日/两日体验套餐。",
        "填写日期、人数、预算、酒店区域与想去地点，生成更顺路的路线方向。",
        f"在司机请求页选择 {d['name']}，把同一份需求通过邮件发给司机确认。",
        "先看网站的初始价格，再等待司机根据实际天数、路线、接送和超时规则回复最终报价。",
    ):
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph("游客问你时，可以这样回答", style="Heading 2")
    add_copy_box(doc, "路线问题", "请先在 WanderMind 选择公开路线或体验套餐，再把日期、人数、酒店和想去地点一起提交。这样我可以更准确地检查路况与每天能否完成。")
    add_copy_box(doc, "价格问题", "网站显示的是初始价格。收到请求后，我会根据用车天数、路线、接送区域和工作时长确认最终报价。")
    add_copy_box(doc, "是否已经预订", "提交表单只是询价与可用性确认；在日期、车辆、路线和最终价格都确认前，还不算完成预订。")


def add_quick_start(doc: Document, d: dict) -> None:
    doc.add_page_break()
    doc.add_paragraph("从这里开始", style="Heading 1")
    doc.add_paragraph("每 3–5 天发布一次即可", style="Heading 2")
    doc.add_paragraph("不用自己重新写文案。按下面顺序操作，再复制后续页面中标有“可复制”的内容。")
    steps = [
        ("保存 7 张图片", "只使用本手册同目录内编号 01–07 的图片；同样的图片也嵌入在本文件里。"),
        ("把专属链接放到个人主页", "Instagram：编辑个人资料 → 链接 → 添加外部链接，然后粘贴下方链接。"),
        ("发布第一条图文", "可用个人旅行帅照、服务照片或风景照做首图；复制现成文案，不要改错专属链接。"),
        ("补充 Story 或 WhatsApp Status", "使用已经准备好的短句；如果平台支持，在第三屏加入 Link sticker。"),
        ("把咨询统一引导到表单", "不要在评论或私信里索取酒店、航班、预算或其他个人信息。"),
    ]
    for idx, (title, body) in enumerate(steps, 1):
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        set_cell_width(table.cell(0, 0), 720)
        set_cell_width(table.cell(0, 1), 8640)
        shade(table.cell(0, 0), GOLD)
        shade(table.cell(0, 1), PAPER)
        for c in table.rows[0].cells:
            cell_margins(c, 110, 120, 110, 120)
        p = table.cell(0, 0).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(idx))
        r.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(255, 255, 255)
        p = table.cell(0, 1).paragraphs[0]
        r = p.add_run(title)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(TEAL)
        p.add_run(f"\n{body}")
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    add_copy_box(doc, f"{d['name']} 的专属链接", d["link"])
    doc.add_paragraph(f"发布前先打开一次，确认页面自动选中 {d['name']}。链接不能换成另一位司机的版本。")


def add_primary_photos(doc: Document, d: dict) -> None:
    doc.add_page_break()
    doc.add_paragraph("第一组图片", style="Heading 1")
    doc.add_paragraph("3 张服务照片：可选个人、服务时刻与车辆", style="Heading 2")
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for cell in row.cells:
            set_cell_width(cell, 4680)
    photos = d["photos"][:3]
    add_photo_card(table.cell(0, 0), d["pack"] / photos[0][0], photos[0][1], photos[0][2], Inches(2.45))
    add_photo_card(table.cell(0, 1), d["pack"] / photos[1][0], photos[1][1], photos[1][2], Inches(2.45))
    add_photo_card(table.cell(1, 0), d["pack"] / photos[2][0], photos[2][1], photos[2][2], Inches(2.45))
    shade(table.cell(1, 1), PALE_TEAL)
    cell_margins(table.cell(1, 1), 180, 180, 180, 180)
    p = table.cell(1, 1).paragraphs[0]
    r = p.add_run("发布顺序")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(TEAL)
    p.add_run("\n\n01 → 02 → 03\n\n这个顺序只是推荐，不是强制。也可以先放自己的旅行帅照或本资料包中的巴厘岛风景。")
    p = table.cell(1, 1).add_paragraph("车辆照片只是参考；没有核实前，不承诺具体车型、座位数、行李容量、价格、时段或一定可用。")
    p.runs[0].font.size = Pt(8.5)
    p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)


def add_feed_copy(doc: Document, d: dict) -> None:
    doc.add_page_break()
    doc.add_paragraph("Instagram · Feed", style="Heading 1")
    doc.add_paragraph("第一条发布：照片 01 → 02 → 03", style="Heading 2")
    caption = (
        "巴厘岛在地图上看起来不大，但舒适的一天仍需要合理、顺路的路线。\n\n"
        f"我是 {d['name']}，巴厘岛本地司机和向导。我主要帮助客人{d['service']}"
        "通过 WanderMind，旅行者可以在一个表单中提交日期、人数、想去的区域和车辆需求。\n\n"
        f"先查看巴厘岛公开路线，选择 {d['name']}，再通过个人主页链接提交请求：\n{d['link']}\n\n"
        "网站显示初始价格。我会在阅读请求后，根据用车天数、路线、接送区域和工作时长确认最终报价。为了保护隐私，请不要在评论中发布酒店、航班或个人资料。\n\n"
        "你最想先安排巴厘岛的哪一天：南部海岸、乌布、日出，还是 Nusa Penida？\n\n"
        "#BaliDriver #BaliItinerary #VisitBali #BaliTravel #WanderMind"
    )
    add_copy_box(doc, "可复制 · 印尼语 Caption 的中文译文", caption)
    caption_en = (
        "巴厘岛在地图上看起来不大，但舒适的一天仍需要合理的路线。\n\n"
        f"我是 {d['name']}，巴厘岛本地司机和向导。旅行者可以通过 WanderMind，一次提交日期、同行人数、偏好区域和车辆需求。\n\n"
        f"查看巴厘岛公开路线，选择 {d['name']}，再通过个人主页链接提交请求：\n{d['link']}\n\n"
        "网站显示初始价格。可用时间、时长、路线和最终报价都会在审核后确认。请不要在评论中发布酒店、航班或个人资料。\n\n"
        "你会先规划巴厘岛的哪一天：南部海岸、乌布、日出，还是 Nusa Penida？\n\n"
        "#BaliDriver #BaliItinerary #VisitBali #BaliTravel #WanderMind"
    )
    add_copy_box(doc, "可复制 · 英文 Caption 的中文译文", caption_en)


def add_channels(doc: Document, d: dict) -> None:
    doc.add_page_break()
    doc.add_paragraph("Story · Facebook · WhatsApp · Reels", style="Heading 1")
    doc.add_paragraph("Instagram Story：3 屏", style="Heading 2")
    story = [
        "第 1 屏：不要从一长串地点开始规划巴厘岛。",
        "第 2 屏：先在 WanderMind 选择公开路线、日期和同行人数。",
        f"第 3 屏：选择 {d['name']}，再提交一份完整请求。\n{d['link']}",
    ]
    for line in story:
        add_copy_box(doc, "可复制", line)
    doc.add_paragraph("操作：Instagram → Your story → 依次选 01/02/03 → 每屏贴一句 → 第 3 屏加 Link sticker → Share。")

    facebook = (
        f"正在准备巴厘岛旅行吗？我是 {d['name']}，巴厘岛本地司机。WanderMind 可以帮助你先选择旅行方向，再通过一个表单提交日期、人数和车辆需求。"
        "网站显示初始价格；我会在阅读请求后根据天数、路线、接送和时长确认最终报价。\n\n"
        f"提交给 {d['name']}：\n{d['link']}\n\n请不要在评论中写酒店、航班或个人信息。"
    )
    add_copy_box(doc, "可复制 · Facebook", facebook)
    doc.add_paragraph("WhatsApp Status", style="Heading 2")
    status_table = doc.add_table(rows=3, cols=2)
    status_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    status_table.autofit = False
    status_rows = [
        ("Status 1", "正在规划巴厘岛？先从顺路的路线开始，而不是一长串地点。"),
        ("Status 2", f"通过 WanderMind，把日期、人数和路线一次提交给 {d['name']}。"),
        ("Status 3", d["link"]),
    ]
    for ri, row in enumerate(status_rows):
        set_cell_width(status_table.cell(ri, 0), 1550)
        set_cell_width(status_table.cell(ri, 1), 7810)
        for cell in status_table.rows[ri].cells:
            cell_margins(cell, 65, 100, 65, 100)
            if ri % 2 == 0:
                shade(cell, PAPER)
        status_table.cell(ri, 0).text = row[0]
        status_table.cell(ri, 1).text = row[1]
        status_table.cell(ri, 0).paragraphs[0].runs[0].bold = True
        status_table.cell(ri, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(TEAL)


def add_reels_and_replies(doc: Document, d: dict) -> None:
    doc.add_page_break()
    doc.add_paragraph("20 秒 Reels 与常用回复", style="Heading 1")
    table = doc.add_table(rows=6, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [1250, 2900, 5210]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            cell_margins(cell, 85, 100, 85, 100)
    headers = ["时间", "画面", "屏幕文字（中文含义）"]
    for i, value in enumerate(headers):
        table.cell(0, i).text = value
        shade(table.cell(0, i), TEAL)
        for run in table.cell(0, i).paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    rows = [
        ("0–3 秒", "巴厘岛照片/视频", "巴厘岛看起来不大，但路线依然重要。"),
        ("3–8 秒", "WanderMind 路线", "先选方向，再添加更多地点。"),
        ("8–13 秒", f"{d['name']} 与客人", f"向 {d['name']} 提交一份完整请求。"),
        ("13–17 秒", "车辆", "先看估价，审核后确认最终价格。"),
        ("17–20 秒", "WanderMind + 链接", "先规划路线，再让它真正落地。"),
    ]
    for ri, row in enumerate(rows, 1):
        for ci, value in enumerate(row):
            table.cell(ri, ci).text = value
            if ri % 2 == 0:
                shade(table.cell(ri, ci), PAPER)
    add_copy_box(doc, "可复制 · Reels Caption", f"巴厘岛的路线如果每天保持顺路，会更容易真正执行。先选路线，再通过个人主页链接向 {d['name']} 提交一份完整请求。#BaliDriver #BaliRoute #WanderMind")

    doc.add_paragraph("有人询问时", style="Heading 2")
    replies = [
        ("询问价格", f"谢谢。网站显示初始价格；最终价格需要根据日期、人数、用车天数、接送和路线核对。请在这里向 {d['name']} 提交请求：\n{d['link']}"),
        ("想在私信发送详情", f"为了避免日期和路线遗漏，请不要在私信中发送旅行资料。请使用 WanderMind 表单并选择 {d['name']}：\n{d['link']}"),
        ("已经提交表单", "谢谢。我会先查看日期、路线和车辆需求，然后确认可用时间与最终报价。"),
        ("要求直接预订", "在可用时间、车辆、时长、路线和最终价格确认前，这份请求还不等于完成预订。"),
    ]
    for label, text in replies:
        add_copy_box(doc, f"可复制 · {label}", text)


def add_rules(doc: Document, d: dict) -> None:
    doc.add_page_break()
    doc.add_paragraph("发布规则", style="Heading 1")
    rules = [
        "不要在帖子或评论中写私人邮箱、WhatsApp 号码或客人资料。",
        "检查请求之前，不承诺具体车辆、容量、时间、最终价格或一定可用。",
        "不要使用“最低价”“保证有车”“即时预订”或未经核验的安全承诺。",
        "只使用本资料包中的图片，不要随意从互联网另找图片替换。",
    ]
    for item in rules:
        doc.add_paragraph(item, style="List Bullet")
    add_callout(doc, "发布前检查", f"链接能打开 · 自动选中 {d['name']} · 图片编号正确 · 没有公开个人资料", PALE_TEAL)
    doc.add_paragraph("图片来源说明", style="Heading 2")
    doc.add_paragraph("照片 01–03 已记录为 user_provided_with_consent；照片 04–07 使用已记录许可和署名的素材。带许可的图片不能写成司机本人拍摄，也不能删除署名。")

    doc.add_paragraph("第二条发布：4 张巴厘岛风景", style="Heading 1")
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for cell in row.cells:
            set_cell_width(cell, 4680)
    for idx, photo in enumerate(d["photos"][3:]):
        cell = table.cell(idx // 2, idx % 2)
        add_photo_card(cell, d["pack"] / photo[0], photo[1], photo[2], Inches(2.08))
    doc.add_paragraph("顺序：04 → 05 → 06 → 07。把四张图作为 carousel；专属链接放在个人主页、Story 或私信中。")


def add_scenic_copy_and_credits(doc: Document, d: dict) -> None:
    if d["name"] == "Dicky":
        doc.add_page_break()
    doc.add_paragraph("第二条风景发布文案", style="Heading 1")
    add_copy_box(doc, "可复制 · 印尼语文案的中文译文", f"{d['scenic_caption']}\n\n{d['utm']}")
    add_copy_box(doc, "可复制 · 英文文案的中文译文", f"{d['scenic_caption']}\n\n{d['utm']}")
    doc.add_paragraph("图片署名 · 不得删除", style="Heading 2")
    credits = doc.add_table(rows=5, cols=3)
    credits.alignment = WD_TABLE_ALIGNMENT.CENTER
    credits.autofit = False
    widths = [900, 3900, 4560]
    for row in credits.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            cell_margins(cell, 90, 110, 90, 110)
    for i, value in enumerate(("编号", "地点", "摄影者与许可")):
        credits.cell(0, i).text = value
        shade(credits.cell(0, i), TEAL)
        for run in credits.cell(0, i).paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    for ri, photo in enumerate(d["photos"][3:], 1):
        credits.cell(ri, 0).text = photo[1].split("·", 1)[0].strip()
        credits.cell(ri, 1).text = photo[1].split("·", 1)[1].strip()
        credits.cell(ri, 2).text = photo[2]
        if ri % 2 == 0:
            for cell in credits.rows[ri].cells:
                shade(cell, PAPER)
    add_callout(doc, "许可提醒", "这些图片不能写成司机本人的摄影作品；WanderMind 只为本推广包保存副本。署名、许可和 ShareAlike 条件继续有效。")


def build_promotion_manual(key: str, d: dict) -> Path:
    doc = Document()
    set_doc_defaults(doc)
    add_cover(doc, d)
    add_quick_start(doc, d)
    add_primary_photos(doc, d)
    add_feed_copy(doc, d)
    add_channels(doc, d)
    add_reels_and_replies(doc, d)
    add_rules(doc, d)
    add_scenic_copy_and_credits(doc, d)
    out = d["pack"] / f"WanderMind_{d['name'].replace(' ', '_')}_Promotion_Chinese_Review.docx"
    doc.save(out)
    return out


def add_field_line(doc: Document, label: str, value: str = "____________________________") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(label)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)
    p.add_run(value)


def add_form_section(doc: Document, number: int, title: str, fields: list[str]) -> None:
    doc.add_paragraph(f"{number}. {title}", style="Heading 2")
    for field in fields:
        add_field_line(doc, f"{field}： ")


def build_rate_form(d: dict) -> Path:
    doc = Document()
    set_doc_defaults(doc)
    add_brand_rule(doc)
    p = doc.add_paragraph("司机报价授权表", style="Title")
    doc.add_paragraph(f"{d['name']} · 中文审阅版", style="Subtitle")
    add_callout(doc, "用途", "让 WanderMind 可以显示透明的路线估价规则，而不是虚构最终价格。最终价格仍由司机根据日期、酒店、路线、人数和工作时长逐单确认。")
    add_callout(doc, "填写要求", "请司机填写金额或写“不适用”。不要提供银行账户、密码、身份证件或其他个人敏感资料。未填字段继续保持未授权状态。", PALE_TEAL)
    add_field_line(doc, "费率生效日期")
    add_field_line(doc, "币种", "IDR（如不是 IDR，请注明）")

    add_form_section(doc, 1, "全天包车", [
        "最多包含小时数", "车辆 + 司机基础价（Rp）", "每位客人每天附加费（Rp）",
        "是否含燃油（是/否）", "是否含停车（是/否）", "是否含过路费（是/否）", "是否含司机餐费（是/否）",
    ])
    add_form_section(doc, 2, "半天包车", ["最多包含小时数", "车辆 + 司机基础价（Rp）", "每位客人每天附加费（Rp）"])
    add_form_section(doc, 3, "超时", ["全天包车超时价（Rp/小时）", "半天包车超时价（Rp/小时）", "计时取整规则（每多少分钟/小时）"])

    doc.add_page_break()
    doc.add_paragraph("接送与跨区域规则", style="Heading 1")
    add_form_section(doc, 4, "机场接送", [
        "机场 → Kuta / Seminyak / Canggu（Rp）", "机场 → Sanur（Rp）", "机场 → Ubud（Rp）",
        "机场 → 其他区域（区域 + 价格）", "航班延误免费等待时长",
    ])
    add_form_section(doc, 5, "携带行李更换酒店", [
        "如包含在当天游览中：免费或附加费", "仅更换酒店时的计算方式", "可接受的行李数量/尺寸限制",
    ])
    add_form_section(doc, 6, "Nusa Penida", [
        f"{d['name']} 仅送到码头，还是可售船票 + 岛上用车套餐",
        "出发码头", "往返快船每人价格（Rp）", "Penida 岛上车辆 + 司机每日价格（Rp）",
        "West / East / Combination 是否不同价", "包含/不包含：门票、停车、餐食、酒店接送",
    ])

    doc.add_page_break()
    doc.add_paragraph("区域附加费与其他规则", style="Heading 1")
    add_form_section(doc, 7, "区域附加费", [
        "Uluwatu / South Bali（Rp）", "Ubud / Central Bali（Rp）", "Kintamani / Batur（Rp）",
        "Karangasem / Amed / East Bali（Rp）", "Munduk / Lovina / North Bali（Rp）", "司机需要异地过夜时每晚费用（Rp）",
    ])
    add_form_section(doc, 8, "其他规则", [
        "舒适乘坐人数 + 行李容量", "儿童座椅是否提供及价格", "取消 / 天气 / no-show 规则",
        "报价有效天数", "税费或其他费用",
    ])
    add_callout(doc, "授权声明", f"我允许 WanderMind 按本表规则，以 {d['name']} 的名义展示估价。每次旅行的最终价格仍由我确认。")
    add_field_line(doc, "同意上述授权（是/否）")
    add_field_line(doc, "姓名", d["name"])
    add_field_line(doc, "确认日期")

    doc.add_paragraph("Dicky 已提供的初始价格", style="Heading 1")
    add_callout(doc, "价格说明", "全天 Rp700,000（最多 10 小时）、半天 Rp500,000（最多 6 小时）、每位客人每天附加 Rp50,000、全天超时 Rp75,000/小时。Dicky 可在查看用车天数和路线后调整；最终金额以司机回复为准。", PAPER)
    doc.add_paragraph("给用户的审阅问题", style="Heading 2")
    for item in (
        "这些字段是否覆盖你和司机谈过的全部收费情况？",
        "是否需要新增夜间接送、节假日、特殊车型或儿童座椅规则？",
        "司机回复后，是否允许 WanderMind 把已确认规则用于路线估价，但仍保留逐单最终确认？",
    ):
        doc.add_paragraph(item, style="List Bullet")

    safe_name = d["name"].replace(" ", "_")
    out = ROOT / "wandermind-studio/operations" / f"WanderMind_{safe_name}_Rate_Authorization_Chinese_Review.docx"
    doc.save(out)
    return out


def add_mobile_image_saving(doc: Document, d: dict) -> None:
    doc.add_page_break()
    doc.add_paragraph("手机保存图片", style="Heading 1")
    add_callout(doc, "最简单的方法", "WanderMind 会同时提供本 Word 文件和一个 ZIP 图片包。Word 里可以直接看图；真正发布时，优先从 ZIP 解压后的 01–07 JPG 原图选择，清晰度更高。", PALE_TEAL)
    doc.add_paragraph("iPhone", style="Heading 2")
    for item in ("在“文件”App 中点开 ZIP，系统会自动生成同名文件夹。", "打开文件夹，点开需要的 JPG。", "点左下角分享按钮 → 选择“存储图像”；随后可在“照片”App 选择发布。"):
        doc.add_paragraph(item, style="List Number")
    doc.add_paragraph("Android", style="Heading 2")
    for item in ("在 Files / My Files / 下载 中点开 ZIP，选择“解压”。", "打开解压后的 images 文件夹，再打开需要的 JPG。", "选择“分享”直接发布，或移动/保存到 Pictures 后再从相册选择。"):
        doc.add_paragraph(item, style="List Number")
    add_callout(doc, "发布节奏", "每 3–5 天发布一次即可。现在先使用本文件中的两组完整素材；以后站长会定期发送新的图片与文案包，继续按同样方式复制发布。")


def add_rate_authorization_to_doc(doc: Document, d: dict) -> None:
    doc.add_page_break()
    doc.add_paragraph("司机报价授权表", style="Title")
    doc.add_paragraph(f"{d['name']} · 与推广指南合并审阅", style="Subtitle")
    if d["name"] == "Dicky":
        price_note = "网站展示的是 Dicky 已口头提供的初始价格。游客提交请求后，Dicky 可根据用车天数、路线、接送区域和工作时长调整；最终金额以司机邮件回复为准。"
    else:
        price_note = "Dicky 提供的初始价格不会自动套用到 Gede Nico。请 Gede Nico 在本表填写自己的初始价格与适用条件；最终金额仍以收到游客请求后的邮件回复为准。"
    add_callout(doc, "价格说明", price_note, PALE_TEAL)
    add_callout(doc, "填写要求", "请填写金额或写“不适用”。不要提供银行账户、密码、身份证件或其他个人敏感资料。未填字段继续保持未授权状态。")
    add_field_line(doc, "费率生效日期")
    add_field_line(doc, "币种", "IDR（如不是 IDR，请注明）")
    add_form_section(doc, 1, "全天包车", ["最多包含小时数", "车辆 + 司机初始价（Rp）", "每位客人每天附加费（Rp）", "是否含燃油（是/否）", "是否含停车（是/否）", "是否含过路费（是/否）", "是否含司机餐费（是/否）"])
    add_form_section(doc, 2, "半天包车", ["最多包含小时数", "车辆 + 司机初始价（Rp）", "每位客人每天附加费（Rp）"])
    add_form_section(doc, 3, "超时", ["全天包车超时价（Rp/小时）", "半天包车超时价（Rp/小时）", "计时取整规则"])
    doc.add_page_break()
    doc.add_paragraph("接送与跨区域规则", style="Heading 1")
    add_form_section(doc, 4, "机场接送", ["机场 → Kuta / Seminyak / Canggu（Rp）", "机场 → Sanur（Rp）", "机场 → Ubud（Rp）", "机场 → 其他区域（区域 + 价格）", "航班延误免费等待时长"])
    add_form_section(doc, 5, "换酒店与行李", ["包含在当天游览时的规则", "仅更换酒店时的价格", "行李数量/尺寸限制"])
    add_form_section(doc, 6, "Nusa Penida", [f"{d['name']} 仅送到码头，还是可安排船票 + 岛上用车", "出发码头", "往返快船每人价格（Rp）", "岛上车辆 + 司机每日价格（Rp）", "West / East / Combination 是否不同价", "门票、停车、餐食、酒店接送包含范围"])
    doc.add_page_break()
    doc.add_paragraph("区域附加费与其他规则", style="Heading 1")
    add_form_section(doc, 7, "区域附加费", ["Uluwatu / South Bali（Rp）", "Ubud / Central Bali（Rp）", "Kintamani / Batur（Rp）", "Karangasem / Amed / East Bali（Rp）", "Munduk / Lovina / North Bali（Rp）", "司机异地过夜每晚费用（Rp）"])
    add_form_section(doc, 8, "其他规则", ["舒适乘坐人数 + 行李容量", "儿童座椅与价格", "取消 / 天气 / no-show 规则", "报价有效天数", "税费或其他费用"])
    add_callout(doc, "授权声明", f"我允许 WanderMind 按本表规则，以 {d['name']} 的名义展示初始估价。每次旅行的最终价格仍由我在收到邮件请求后确认。")
    add_field_line(doc, "同意上述授权（是/否）")
    add_field_line(doc, "姓名", d["name"])
    add_field_line(doc, "确认日期")
    if d["name"] == "Dicky":
        doc.add_paragraph("已知初始价格（请 Dicky 书面确认适用条件）", style="Heading 1")
        add_callout(doc, "初始价格", "全天 Rp700,000（最多 10 小时）、半天 Rp500,000（最多 6 小时）、每位客人每天附加 Rp50,000、全天超时 Rp75,000/小时。Dicky 可在回复邮件时根据天数和路线调整，最终金额以司机回复为准。", PAPER)
    else:
        doc.add_paragraph("Gede Nico 初始价格待书面填写", style="Heading 1")
        add_callout(doc, "不要套用", "Dicky 的初始价格仅属于 Dicky，不能自动写成 Gede Nico 的报价。请 Gede Nico 完成本表后，再把其规则接入网站。", PAPER)


def build_complete_guide(key: str, d: dict) -> Path:
    doc = Document()
    set_doc_defaults(doc)
    add_cover(doc, d)
    add_site_orientation(doc, d)
    add_quick_start(doc, d)
    add_primary_photos(doc, d)
    add_feed_copy(doc, d)
    add_channels(doc, d)
    add_reels_and_replies(doc, d)
    add_rules(doc, d)
    add_scenic_copy_and_credits(doc, d)
    add_mobile_image_saving(doc, d)
    add_rate_authorization_to_doc(doc, d)
    safe_name = d["name"].replace(" ", "_")
    out = d["pack"] / f"WanderMind_{safe_name}_Complete_Guide_Chinese_Review.docx"
    doc.save(out)
    return out


INDONESIAN_CONTENT = {
    "dicky": {
        "service": "antar-jemput bandara, perjalanan privat satu hari, dan rute yang disusun sesuai arah perjalanan tamu.",
        "service_en": "provide airport transfers, private day trips, and routes arranged around each guest's travel direction.",
        "photos": [
            ("01-Dicky-profile.jpg", "01 · Foto pribadi", "Opsional. Jika tidak ingin memakai foto pribadi, mulai dengan foto layanan atau pemandangan."),
            ("02-Dicky-guest-moment.jpg", "02 · Momen bersama tamu", "Tampilkan pengalaman layanan yang nyata; jangan menulis identitas tamu."),
            ("03-Dicky-vehicle.jpg", "03 · Referensi kendaraan", "Hanya referensi visual. Kendaraan, kursi, dan kapasitas bagasi harus dikonfirmasi per permintaan."),
            ("04-Dicky-Melasti-Beach.jpg", "04 · Pantai Melasti", "Dare2Leap · CC BY-SA 4.0"),
            ("05-Dicky-Jimbaran-Sunset.jpg", "05 · Teluk Jimbaran", "Simon Sees · CC BY 2.0"),
            ("06-Dicky-Broken-Beach.jpg", "06 · Broken Beach", "Aaron Rentfrew · CC BY-SA 4.0"),
            ("07-Dicky-Tegallalang.jpg", "07 · Tegallalang", "Philip Nalangan · CC BY 4.0"),
        ],
        "scenic_caption": "Empat suasana Bali dalam satu perjalanan: pesisir selatan Melasti, matahari terbenam di Jimbaran, lengkungan alami Broken Beach, dan sawah bertingkat Tegallalang. Simpan inspirasinya, lalu buka tautan WanderMind saya untuk mengirim rencana perjalanan. Rute, tanggal, kendaraan, durasi, dan harga final akan dikonfirmasi setelah permintaan diterima.",
        "scenic_caption_en": "Four Bali moods in one journey: the southern coast at Melasti, sunset in Jimbaran, the natural arch at Broken Beach, and the layered rice fields of Tegallalang. Save the inspiration, then open my WanderMind link to send your travel plan. Route, dates, vehicle, duration and final price will be confirmed after the request is received.",
    },
    "gede": {
        "service": "menghubungkan budaya, makanan, aktivitas, dan pengalaman lokal Bali dalam hari yang lebih rapi dan mudah dijalankan.",
        "service_en": "connect Bali's culture, food, activities and local experiences into a smoother day that is easier to follow.",
        "photos": [
            ("01-Gede-profile.jpg", "01 · Foto pribadi", "Opsional. Jika tidak ingin memakai foto pribadi, mulai dengan foto layanan atau pemandangan."),
            ("02-Gede-guest-moment.jpg", "02 · Momen bersama tamu", "Tampilkan pengalaman layanan yang nyata; jangan menulis identitas tamu."),
            ("03-Gede-vehicle.jpg", "03 · Referensi kendaraan", "Hanya referensi visual. Kendaraan, kursi, dan kapasitas bagasi harus dikonfirmasi per permintaan."),
            ("04-Gede-Campuhan-Ridge.jpg", "04 · Campuhan Ridge Walk", "Artem Beliaikin · CC0 1.0"),
            ("05-Gede-Jatiluwih.jpg", "05 · Jatiluwih", "Jorge Franganillo · CC BY 2.0"),
            ("06-Gede-Tirta-Gangga.jpg", "06 · Tirta Gangga", "Bair175 · CC BY-SA 3.0"),
            ("07-Gede-Seminyak-Sunset.jpg", "07 · Pantai Seminyak", "Christophe95 · CC BY-SA 4.0"),
        ],
        "scenic_caption": "Empat suasana Bali dalam satu perjalanan: pagi hijau di Campuhan, sawah bertingkat Jatiluwih, taman air Tirta Gangga, dan matahari terbenam di Seminyak. Simpan inspirasinya, lalu buka tautan WanderMind saya untuk mengirim rencana perjalanan. Rute, tanggal, kendaraan, durasi, dan harga final akan dikonfirmasi setelah permintaan diterima.",
        "scenic_caption_en": "Four Bali moods in one journey: a green morning at Campuhan, the layered rice fields of Jatiluwih, the water gardens of Tirta Gangga, and sunset in Seminyak. Save the inspiration, then open my WanderMind link to send your travel plan. Route, dates, vehicle, duration and final price will be confirmed after the request is received.",
    },
}


def set_doc_defaults_id(doc: Document) -> None:
    """Keep the approved WanderMind teal/gold system, with Indonesian running labels."""
    set_doc_defaults(doc)
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "WanderMind Studio · Panduan kerja sama"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(TEAL)
    footer = section.footer.paragraphs[0]
    footer.text = "Panduan kerja sama · Periksa tautan sebelum menerbitkan"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)


def add_id_cover(doc: Document, d: dict, c: dict) -> None:
    add_brand_rule(doc)
    p = doc.add_paragraph("WanderMind Studio", style="Subtitle")
    p.paragraph_format.space_before = Pt(18)
    doc.add_paragraph("Panduan Lengkap Kerja Sama & Promosi", style="Title")
    p = doc.add_paragraph(d["name"], style="Title")
    p.runs[0].font.color.rgb = RGBColor.from_string(GOLD)
    doc.add_paragraph("Pengenalan situs · Panduan ponsel · Materi publikasi · Otorisasi tarif", style="Subtitle")

    hero = d["pack"] / c["photos"][0][0]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(image_stream(hero), width=Inches(4.35))
    add_callout(doc, "Cara memakai panduan ini", "Gunakan tautan, foto, dan teks yang berada di bagian pengemudi Anda sendiri. Jangan menukar tautan atau materi dengan paket pengemudi lain.", PALE_TEAL)
    doc.add_paragraph("7 gambar siap pakai · 2 materi publikasi · 1 tautan khusus", style="Subtitle")


def add_id_site_orientation(doc: Document, d: dict) -> None:
    doc.add_page_break()
    doc.add_paragraph("Kenali WanderMind", style="Heading 1")
    add_callout(doc, "Tujuan kami", "WanderMind membantu wisatawan memahami tempat yang nyata, kekuatan rute, dan waktu perjalanan sebelum mereka meminta konfirmasi kepada pengemudi lokal. Foto inspirasi tidak dianggap sebagai jaminan kondisi di lapangan, dan tarif yang belum dikonfirmasi tidak ditulis sebagai harga final.", PALE_TEAL)
    doc.add_paragraph("Apa yang dapat dilakukan wisatawan di situs", style="Heading 2")
    for item in (
        "Melihat rute Bali, foto tempat, dan paket pengalaman satu atau dua hari yang dapat disesuaikan.",
        "Mengisi tanggal, jumlah tamu, anggaran, area hotel, dan tempat yang ingin dikunjungi.",
        f"Memilih {d['name']} di halaman permintaan pengemudi, lalu mengirim detail melalui formulir untuk diperiksa.",
        "Melihat tarif awal bila tersedia, lalu menunggu balasan pengemudi untuk waktu, rute, kendaraan, dan harga final.",
    ):
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph("Jawaban singkat untuk wisatawan", style="Heading 2")
    add_copy_box(doc, "Pertanyaan tentang rute", "Silakan pilih rute atau paket pengalaman di WanderMind, lalu kirim tanggal, jumlah tamu, hotel, dan tempat yang ingin dikunjungi. Dengan begitu saya dapat memeriksa apakah rute tersebut nyaman dijalankan.")
    add_copy_box(doc, "Pertanyaan tentang harga", "Harga di situs adalah tarif awal. Setelah menerima permintaan, saya akan memeriksa jumlah hari, rute, area penjemputan, dan durasi sebelum mengonfirmasi harga final.")
    add_copy_box(doc, "Apakah sudah dipesan?", "Pengiriman formulir berarti permintaan harga dan ketersediaan. Ini belum menjadi pemesanan sampai tanggal, kendaraan, rute, durasi, dan harga final dikonfirmasi.")


def add_id_quick_start(doc: Document, d: dict) -> None:
    doc.add_page_break()
    doc.add_paragraph("Mulai di sini", style="Heading 1")
    doc.add_paragraph("Publikasikan setiap 3–5 hari", style="Heading 2")
    doc.add_paragraph("Tidak perlu menulis ulang. Ikuti langkah singkat ini, lalu salin materi pada halaman berikutnya.")
    steps = [
        ("Simpan 7 gambar", "Gunakan gambar bernomor 01–07 dalam folder ini. Gambar yang sama juga sudah dimasukkan ke dalam dokumen."),
        ("Pasang tautan khusus di bio", "Instagram: Edit profile → Links → Add external link, lalu tempel tautan khusus di bawah."),
        ("Terbitkan materi pertama", "Foto pribadi tidak wajib. Gunakan foto perjalanan yang Anda sukai, foto layanan, atau pemandangan dari paket ini."),
        ("Tambahkan Story atau WhatsApp Status", "Gunakan kalimat yang sudah disiapkan. Jika tersedia, letakkan Link sticker pada layar terakhir."),
        ("Arahkan pertanyaan ke formulir", "Jangan meminta hotel, penerbangan, anggaran, atau data pribadi di kolom komentar atau pesan publik."),
    ]
    for idx, (title, body) in enumerate(steps, 1):
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        set_cell_width(table.cell(0, 0), 720)
        set_cell_width(table.cell(0, 1), 8640)
        shade(table.cell(0, 0), GOLD)
        shade(table.cell(0, 1), PAPER)
        for cell in table.rows[0].cells:
            cell_margins(cell, 110, 120, 110, 120)
        p = table.cell(0, 0).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(idx))
        r.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(255, 255, 255)
        p = table.cell(0, 1).paragraphs[0]
        r = p.add_run(title)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(TEAL)
        p.add_run(f"\n{body}")
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    add_copy_box(doc, f"Tautan khusus {d['name']}", d["link"])
    doc.add_paragraph(f"Buka tautan sekali sebelum menerbitkan dan pastikan halaman otomatis memilih {d['name']}. Jangan mengganti tautan dengan versi pengemudi lain.")


def add_id_primary_photos(doc: Document, d: dict, c: dict) -> None:
    doc.add_page_break()
    doc.add_paragraph("Materi gambar pertama", style="Heading 1")
    doc.add_paragraph("3 foto layanan: pribadi, momen layanan, dan kendaraan", style="Heading 2")
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for cell in row.cells:
            set_cell_width(cell, 4680)
    photos = c["photos"][:3]
    add_photo_card(table.cell(0, 0), d["pack"] / photos[0][0], photos[0][1], photos[0][2], Inches(2.45))
    add_photo_card(table.cell(0, 1), d["pack"] / photos[1][0], photos[1][1], photos[1][2], Inches(2.45))
    add_photo_card(table.cell(1, 0), d["pack"] / photos[2][0], photos[2][1], photos[2][2], Inches(2.45))
    shade(table.cell(1, 1), PALE_TEAL)
    cell_margins(table.cell(1, 1), 180, 180, 180, 180)
    p = table.cell(1, 1).paragraphs[0]
    r = p.add_run("Urutan yang disarankan")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(TEAL)
    p.add_run("\n\n01 → 02 → 03\n\nUrutan ini hanya saran, bukan kewajiban. Anda boleh memakai foto perjalanan yang bagus atau foto pemandangan dari paket ini sebagai gambar pertama.")
    p = table.cell(1, 1).add_paragraph("Foto kendaraan hanya sebagai referensi. Sebelum diperiksa, jangan menjanjikan model kendaraan, jumlah kursi, kapasitas bagasi, durasi, harga, atau ketersediaan tertentu.")
    p.runs[0].font.size = Pt(8.5)
    p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)


def add_id_feed_copy(doc: Document, d: dict, c: dict) -> None:
    doc.add_page_break()
    doc.add_paragraph("Instagram · Feed", style="Heading 1")
    doc.add_paragraph("Publikasi pertama: foto 01 → 02 → 03", style="Heading 2")
    caption = (
        "Bali mungkin terlihat dekat di peta, tetapi hari yang nyaman tetap membutuhkan rute yang masuk akal.\n\n"
        f"Saya {d['name']}, pengemudi lokal di Bali. Saya membantu tamu {c['service']}\n\n"
        "Dengan WanderMind, wisatawan dapat mengirim tanggal, jumlah tamu, area yang ingin dikunjungi, dan kebutuhan kendaraan dalam satu formulir.\n\n"
        f"Lihat rute Bali di situs, pilih {d['name']}, lalu kirim permintaan melalui tautan di bio:\n{d['link']}\n\n"
        "Situs menampilkan tarif awal. Setelah membaca permintaan, saya akan memeriksa tanggal, jumlah hari, rute, area penjemputan, dan durasi sebelum memberi harga final. Demi privasi, jangan tulis hotel, penerbangan, atau data pribadi di komentar.\n\n"
        "Kamu ingin mulai dari pantai selatan, Ubud, matahari terbit, atau Nusa Penida?\n\n"
        "#BaliDriver #BaliItinerary #VisitBali #BaliTravel #WanderMind"
    )
    add_copy_box(doc, "SALIN · Caption Bahasa Indonesia", caption)
    caption_en = (
        "Bali may look close on a map, but a comfortable day still needs a sensible route.\n\n"
        f"I'm {d['name']}, a local driver in Bali. I help guests {c['service_en']}\n\n"
        "With WanderMind, travellers can send dates, group size, preferred areas and vehicle needs in one request. Explore the Bali routes, choose my name, and use the link in my bio:\n"
        f"{d['link']}\n\n"
        "The site shows an initial rate. I confirm availability and the final price after reviewing the request. Please keep hotel, flight and personal details out of public comments.\n\n"
        "#BaliDriver #BaliItinerary #VisitBali #BaliTravel #WanderMind"
    )
    add_copy_box(doc, "SALIN · English caption (opsional)", caption_en)


def add_id_channels(doc: Document, d: dict) -> None:
    doc.add_page_break()
    doc.add_paragraph("Story · Facebook · WhatsApp · Reels", style="Heading 1")
    doc.add_paragraph("Instagram Story: 3 layar", style="Heading 2")
    story = [
        "Layar 1: Jangan mulai merencanakan Bali dari daftar tempat yang panjang.",
        "Layar 2: Pilih dulu arah rute, tanggal, dan jumlah tamu di WanderMind.",
        f"Layar 3: Pilih {d['name']}, lalu kirim satu permintaan lengkap.\n{d['link']}",
    ]
    for line in story:
        add_copy_box(doc, "SALIN", line)
    doc.add_paragraph("Cara posting: Instagram → Your story → pilih foto 01/02/03 secara berurutan → tempel satu kalimat di setiap layar → tambahkan Link sticker pada layar 3 → Share.")

    facebook = (
        f"Sedang menyiapkan perjalanan ke Bali? Saya {d['name']}, pengemudi lokal di Bali. WanderMind membantu wisatawan memilih arah perjalanan terlebih dahulu, lalu mengirim tanggal, jumlah tamu, dan kebutuhan kendaraan melalui satu formulir. "
        "Situs menampilkan tarif awal; saya akan memeriksa jumlah hari, rute, area penjemputan, dan durasi sebelum mengonfirmasi harga final.\n\n"
        f"Kirim permintaan kepada {d['name']}:\n{d['link']}\n\nJangan tulis hotel, penerbangan, atau data pribadi di komentar."
    )
    add_copy_box(doc, "SALIN · Facebook", facebook)
    doc.add_paragraph("WhatsApp Status", style="Heading 2")
    status_table = doc.add_table(rows=3, cols=2)
    status_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    status_table.autofit = False
    status_rows = [
        ("Status 1", "Merencanakan Bali? Mulai dari rute yang masuk akal, bukan daftar tempat yang panjang."),
        ("Status 2", f"Kirim tanggal, jumlah tamu, dan rute melalui WanderMind kepada {d['name']}."),
        ("Status 3", d["link"]),
    ]
    for ri, row in enumerate(status_rows):
        set_cell_width(status_table.cell(ri, 0), 1550)
        set_cell_width(status_table.cell(ri, 1), 7810)
        for cell in status_table.rows[ri].cells:
            cell_margins(cell, 65, 100, 65, 100)
            if ri % 2 == 0:
                shade(cell, PAPER)
        status_table.cell(ri, 0).text = row[0]
        status_table.cell(ri, 1).text = row[1]
        status_table.cell(ri, 0).paragraphs[0].runs[0].bold = True
        status_table.cell(ri, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(TEAL)


def add_id_reels_and_replies(doc: Document, d: dict) -> None:
    doc.add_page_break()
    doc.add_paragraph("Reels 20 detik dan balasan siap pakai", style="Heading 1")
    table = doc.add_table(rows=6, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [1250, 2900, 5210]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            cell_margins(cell, 85, 100, 85, 100)
    headers = ["Waktu", "Visual", "Teks layar"]
    for i, value in enumerate(headers):
        table.cell(0, i).text = value
        shade(table.cell(0, i), TEAL)
        for run in table.cell(0, i).paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    rows = [
        ("0–3 detik", "Foto/video Bali", "Bali terlihat dekat, tetapi rute tetap penting."),
        ("3–8 detik", "Rute WanderMind", "Pilih arah dulu, lalu tambahkan tempat yang sesuai."),
        ("8–13 detik", f"{d['name']} bersama tamu", f"Kirim satu permintaan lengkap kepada {d['name']}."),
        ("13–17 detik", "Kendaraan", "Lihat perkiraan, lalu tunggu konfirmasi harga final."),
        ("17–20 detik", "WanderMind + tautan", "Rencanakan rute, lalu wujudkan di lapangan."),
    ]
    for ri, row in enumerate(rows, 1):
        for ci, value in enumerate(row):
            table.cell(ri, ci).text = value
            if ri % 2 == 0:
                shade(table.cell(ri, ci), PAPER)
    add_copy_box(doc, "SALIN · Caption Reels", f"Rute Bali yang tersusun rapi lebih mudah dijalankan. Pilih rute, lalu kirim permintaan lengkap kepada {d['name']} melalui tautan di bio. #BaliDriver #BaliRoute #WanderMind")

    doc.add_paragraph("Jika ada yang bertanya", style="Heading 2")
    replies = [
        ("Tanya harga", f"Terima kasih. Situs menampilkan tarif awal; harga final perlu disesuaikan dengan tanggal, jumlah tamu, hari penggunaan kendaraan, penjemputan, dan rute. Kirim permintaan kepada {d['name']} di sini:\n{d['link']}"),
        ("Ingin mengirim detail lewat pesan pribadi", f"Agar tanggal dan rute tidak tertinggal, silakan gunakan formulir WanderMind dan pilih {d['name']}. Jangan kirim data perjalanan di ruang publik:\n{d['link']}"),
        ("Sudah mengirim formulir", "Terima kasih. Saya akan memeriksa tanggal, rute, dan kebutuhan kendaraan, lalu membalas ketersediaan dan harga final."),
        ("Ingin langsung booking", "Permintaan belum menjadi booking sampai waktu, kendaraan, durasi, rute, dan harga final dikonfirmasi."),
    ]
    for label, text in replies:
        add_copy_box(doc, f"SALIN · {label}", text)


def add_id_rules_and_scenic_photos(doc: Document, d: dict, c: dict) -> None:
    doc.add_page_break()
    doc.add_paragraph("Aturan publikasi", style="Heading 1")
    rules = [
        "Jangan menulis email pribadi, nomor WhatsApp, atau data tamu di posting dan komentar.",
        "Sebelum memeriksa permintaan, jangan menjanjikan kendaraan, kapasitas, waktu, harga final, atau ketersediaan.",
        "Jangan memakai klaim “harga termurah”, “pasti tersedia”, “booking instan”, atau jaminan keselamatan yang belum diverifikasi.",
        "Gunakan gambar dalam paket ini dan jangan mengambil gambar acak dari internet untuk menggantinya.",
    ]
    for item in rules:
        doc.add_paragraph(item, style="List Bullet")
    add_callout(doc, "Periksa sebelum posting", f"Tautan bisa dibuka · {d['name']} otomatis terpilih · nomor gambar benar · tidak ada data pribadi", PALE_TEAL)
    doc.add_paragraph("Catatan sumber gambar", style="Heading 2")
    doc.add_paragraph("Foto 01–03 dicatat sebagai user_provided_with_consent. Foto 04–07 memakai gambar berlisensi dengan kredit yang tercantum di bagian akhir. Gambar berlisensi tidak boleh ditulis sebagai hasil foto pribadi dan kreditnya tidak boleh dihapus.")

    doc.add_paragraph("Materi publikasi kedua: 4 pemandangan Bali · 04 → 05 → 06 → 07", style="Heading 1")
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for cell in row.cells:
            set_cell_width(cell, 4680)
    scenic_width = Inches(1.8 if d["name"] == "Gede Nico" else 2.08)
    for idx, photo in enumerate(c["photos"][3:]):
        cell = table.cell(idx // 2, idx % 2)
        add_photo_card(cell, d["pack"] / photo[0], photo[1], photo[2], scenic_width)


def add_id_scenic_copy_and_credits(doc: Document, d: dict, c: dict) -> None:
    doc.add_page_break()
    doc.add_paragraph("Caption publikasi pemandangan", style="Heading 1")
    add_copy_box(doc, "SALIN · Caption Bahasa Indonesia", f"{c['scenic_caption']}\n\n{d['utm']}")
    add_copy_box(doc, "SALIN · English caption (opsional)", f"{c['scenic_caption_en']}\n\n{d['utm']}")
    doc.add_paragraph("Kredit gambar · jangan dihapus", style="Heading 2")
    credits = doc.add_table(rows=5, cols=3)
    credits.alignment = WD_TABLE_ALIGNMENT.CENTER
    credits.autofit = False
    widths = [900, 3900, 4560]
    for row in credits.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            cell_margins(cell, 90, 110, 90, 110)
    for i, value in enumerate(("No.", "Tempat", "Fotografer dan lisensi")):
        credits.cell(0, i).text = value
        shade(credits.cell(0, i), TEAL)
        for run in credits.cell(0, i).paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    for ri, photo in enumerate(c["photos"][3:], 1):
        credits.cell(ri, 0).text = photo[1].split("·", 1)[0].strip()
        credits.cell(ri, 1).text = photo[1].split("·", 1)[1].strip()
        credits.cell(ri, 2).text = photo[2]
        if ri % 2 == 0:
            for cell in credits.rows[ri].cells:
                shade(cell, PAPER)
    add_callout(doc, "Pengingat lisensi", "Gambar ini tidak boleh ditulis sebagai foto pribadi pengemudi. WanderMind menyimpan salinan hanya untuk paket promosi ini; kredit dan ketentuan ShareAlike tetap berlaku.")


def add_id_mobile_saving(doc: Document) -> None:
    doc.add_page_break()
    doc.add_paragraph("Cara menyimpan gambar dari ponsel", style="Heading 1")
    add_callout(doc, "Cara paling mudah", "Word dipakai untuk membaca panduan, sedangkan ZIP berisi file JPG yang lebih mudah dipilih ketika akan posting. Unduh ZIP, ekstrak, lalu simpan JPG yang diperlukan ke galeri.", PALE_TEAL)
    doc.add_paragraph("iPhone", style="Heading 2")
    for item in (
        "Buka aplikasi Files, lalu ketuk file ZIP. iPhone akan membuat folder dengan nama yang sama.",
        "Buka folder tersebut dan ketuk JPG yang ingin dipakai.",
        "Ketuk tombol Share di kiri bawah → Save Image. Gambar akan muncul di aplikasi Photos dan dapat dipilih saat posting.",
    ):
        doc.add_paragraph(item, style="List Number")
    doc.add_paragraph("Android", style="Heading 2")
    for item in (
        "Buka Files / My Files / Downloads, ketuk file ZIP, lalu pilih Extract atau Unzip.",
        "Buka folder images hasil ekstraksi, lalu ketuk JPG yang ingin dipakai.",
        "Pilih Share untuk langsung mengirimnya, atau pindahkan ke folder Pictures agar muncul di galeri.",
    ):
        doc.add_paragraph(item, style="List Number")
    add_callout(doc, "Ritme publikasi", "Publikasikan setiap 3–5 hari. Gunakan dua materi lengkap dalam paket ini terlebih dahulu; pemilik situs akan mengirim gambar dan teks baru secara berkala dengan cara yang sama.")


def add_id_field_line(doc: Document, label: str, value: str = "____________________________") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(label)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)
    p.add_run(value)


def add_id_form_section(doc: Document, number: int, title: str, fields: list[str]) -> None:
    doc.add_paragraph(f"{number}. {title}", style="Heading 2")
    for field in fields:
        add_id_field_line(doc, f"{field}: ")


def add_id_rate_authorization(doc: Document, d: dict) -> None:
    doc.add_page_break()
    doc.add_paragraph("Otorisasi tarif pengemudi", style="Title")
    doc.add_paragraph(f"{d['name']} · Formulir konfirmasi", style="Subtitle")
    if d["name"] == "Dicky":
        price_note = "Harga di bawah adalah tarif awal yang disampaikan langsung oleh Dicky, bukan harga final. Setelah menerima permintaan, Dicky dapat menyesuaikannya berdasarkan jumlah hari, rute, area penjemputan, dan durasi; harga final mengikuti balasan email pengemudi."
    else:
        price_note = "Tarif awal Gede Nico harus diisi dan dikonfirmasi sendiri. Jangan memakai tarif pengemudi lain. Setelah menerima permintaan, harga final mengikuti balasan email pengemudi."
    add_callout(doc, "Keterangan harga", price_note, PALE_TEAL)
    add_callout(doc, "Cara mengisi", "Isi nominal atau tulis “tidak berlaku”. Jangan menulis rekening bank, kata sandi, nomor identitas, atau data pribadi sensitif. Kolom kosong berarti belum diberi otorisasi.")
    add_id_field_line(doc, "Tanggal mulai tarif: ")
    add_id_field_line(doc, "Mata uang: ", "IDR (jika berbeda, tulis di sini)")

    add_id_form_section(doc, 1, "Perjalanan satu hari", [
        "Maksimal jam layanan", "Tarif awal kendaraan + pengemudi (Rp)", "Tambahan per tamu untuk setiap hari (Rp)",
        "Termasuk bahan bakar (ya/tidak)", "Termasuk parkir (ya/tidak)", "Termasuk tol (ya/tidak)", "Termasuk makan pengemudi (ya/tidak)",
    ])
    add_id_form_section(doc, 2, "Perjalanan setengah hari", ["Maksimal jam layanan", "Tarif awal kendaraan + pengemudi (Rp)", "Tambahan per tamu untuk setiap hari (Rp)"])
    add_id_form_section(doc, 3, "Lembur", ["Tarif lembur perjalanan satu hari (Rp/jam)", "Tarif lembur setengah hari (Rp/jam)", "Aturan pembulatan waktu"])

    doc.add_page_break()
    doc.add_paragraph("Penjemputan dan aturan lintas wilayah", style="Heading 1")
    add_id_form_section(doc, 4, "Antar-jemput bandara", [
        "Bandara → Kuta / Seminyak / Canggu (Rp)", "Bandara → Sanur (Rp)", "Bandara → Ubud (Rp)",
        "Bandara → wilayah lain (wilayah + harga)", "Waktu tunggu gratis jika penerbangan terlambat",
    ])
    add_id_form_section(doc, 5, "Pindah hotel dan bagasi", [
        "Jika dilakukan saat tur: gratis atau tambahan", "Jika hanya pindah hotel: cara menghitung", "Batas jumlah / ukuran bagasi",
    ])
    add_id_form_section(doc, 6, "Nusa Penida", [
        f"{d['name']} hanya mengantar ke pelabuhan, atau dapat mengatur tiket kapal + kendaraan di pulau",
        "Pelabuhan keberangkatan", "Harga kapal pulang-pergi per orang (Rp)", "Harga kendaraan + pengemudi per hari di pulau (Rp)",
        "Apakah West / East / Combination berbeda harga", "Yang termasuk / tidak termasuk: tiket masuk, parkir, makan, antar-jemput hotel",
    ])

    doc.add_page_break()
    doc.add_paragraph("Tambahan wilayah dan aturan lain", style="Heading 1")
    add_id_form_section(doc, 7, "Tambahan wilayah", [
        "Uluwatu / South Bali (Rp)", "Ubud / Central Bali (Rp)", "Kintamani / Batur (Rp)",
        "Karangasem / Amed / East Bali (Rp)", "Munduk / Lovina / North Bali (Rp)", "Biaya menginap pengemudi di luar area per malam (Rp)",
    ])
    add_id_form_section(doc, 8, "Aturan lain", [
        "Jumlah penumpang nyaman + kapasitas bagasi", "Kursi anak dan harga", "Aturan pembatalan / cuaca / no-show",
        "Masa berlaku penawaran", "Pajak atau biaya lain",
    ])
    add_callout(doc, "Pernyataan otorisasi", f"Saya mengizinkan WanderMind menampilkan perkiraan awal atas nama {d['name']} berdasarkan aturan dalam formulir ini. Harga final setiap perjalanan tetap harus saya konfirmasi setelah menerima permintaan melalui email.")
    add_id_field_line(doc, "Menyetujui otorisasi ini (ya/tidak): ")
    add_id_field_line(doc, "Nama: ", d["name"])
    add_id_field_line(doc, "Tanggal konfirmasi: ")

    if d["name"] == "Dicky":
        doc.add_paragraph("Tarif awal yang telah disampaikan Dicky", style="Heading 1")
        add_callout(doc, "Tarif awal, bukan harga final", "Full day: Rp700.000, maksimal 10 jam. Half day: Rp500.000, maksimal 6 jam. Tambahan: Rp50.000 per tamu untuk setiap hari yang dipilih. Lembur full day: Rp75.000 per jam. Tarif dapat disesuaikan setelah Dicky melihat jumlah hari dan rute; harga final mengikuti balasan email Dicky.", PAPER)
    else:
        doc.add_paragraph("Tarif awal Gede Nico menunggu konfirmasi tertulis", style="Heading 1")
        add_callout(doc, "Belum diisi", "Silakan lengkapi nominal dan syarat penggunaan pada formulir ini. Tarif yang belum ditulis tidak boleh ditampilkan sebagai harga situs atau digunakan untuk menghitung penawaran.", PAPER)


def build_indonesian_guide(key: str, d: dict) -> Path:
    c = INDONESIAN_CONTENT[key]
    doc = Document()
    set_doc_defaults_id(doc)
    add_id_cover(doc, d, c)
    add_id_site_orientation(doc, d)
    add_id_quick_start(doc, d)
    add_id_primary_photos(doc, d, c)
    add_id_feed_copy(doc, d, c)
    add_id_channels(doc, d)
    add_id_reels_and_replies(doc, d)
    add_id_rules_and_scenic_photos(doc, d, c)
    add_id_scenic_copy_and_credits(doc, d, c)
    add_id_mobile_saving(doc)
    add_id_rate_authorization(doc, d)
    safe_name = d["name"].replace(" ", "_")
    out = d["pack"] / f"WanderMind_{safe_name}_Panduan_Lengkap_ID.docx"
    doc.save(out)
    return out


def build_indonesian_mobile_pack(key: str, d: dict, guide: Path) -> Path:
    c = INDONESIAN_CONTENT[key]
    safe_name = d["name"].replace(" ", "_")
    out = d["pack"] / f"WanderMind_{safe_name}_Panduan_Lengkap_ID_Mobile_Pack.zip"
    saving_guide = (
        "PANDUAN MENYIMPAN GAMBAR DARI PONSEL\n\n"
        "iPhone\n"
        "1. Buka aplikasi Files dan ketuk file ZIP. iPhone akan membuat folder dengan nama yang sama.\n"
        "2. Buka folder tersebut dan ketuk JPG yang ingin dipakai.\n"
        "3. Ketuk Share → Save Image. Gambar akan muncul di Photos.\n\n"
        "Android\n"
        "1. Buka Files / My Files / Downloads, ketuk ZIP, lalu pilih Extract atau Unzip.\n"
        "2. Buka folder images hasil ekstraksi dan pilih JPG.\n"
        "3. Pilih Share untuk langsung memakainya, atau pindahkan ke Pictures agar muncul di galeri.\n\n"
        "Publikasikan setiap 3–5 hari. Gunakan teks dan tautan khusus milik pengemudi ini."
    )
    with ZipFile(out, "w", ZIP_DEFLATED) as archive:
        archive.write(guide, arcname=guide.name)
        for filename, _, _ in c["photos"]:
            image = d["pack"] / filename
            if not image.exists():
                raise FileNotFoundError(image)
            archive.write(image, arcname=f"images/{filename}")
        archive.writestr("Panduan_Simpan_Gambar_ID.txt", saving_guide)
        archive.writestr("README_ID.txt", f"Paket resmi WanderMind untuk {d['name']}. Gunakan hanya tautan dan gambar yang ada di paket ini.\n{d['link']}\n")
    return out


def main() -> None:
    outputs = []
    for key, driver in DRIVERS.items():
        outputs.append(build_complete_guide(key, driver))
        guide = build_indonesian_guide(key, driver)
        pack = build_indonesian_mobile_pack(key, driver, guide)
        outputs.extend((guide, pack))
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
